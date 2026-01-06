from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

# Create a new Document
doc = Document()

# Set up styles
styles = doc.styles

# Title style
title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = 'Arial'
title_style.font.size = Pt(24)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)
title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_style.paragraph_format.space_after = Pt(12)

# Heading 1 style
heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
heading1_style.font.name = 'Arial'
heading1_style.font.size = Pt(18)
heading1_style.font.bold = True
heading1_style.font.color.rgb = RGBColor(0, 0, 0)
heading1_style.paragraph_format.space_before = Pt(12)
heading1_style.paragraph_format.space_after = Pt(6)

# Heading 2 style
heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
heading2_style.font.name = 'Arial'
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True
heading2_style.font.color.rgb = RGBColor(0, 0, 0)
heading2_style.paragraph_format.space_before = Pt(12)
heading2_style.paragraph_format.space_after = Pt(6)

# Body text style
body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
body_style.font.name = 'Arial'
body_style.font.size = Pt(11)
body_style.paragraph_format.space_after = Pt(6)
body_style.paragraph_format.line_spacing = 1.5

# COVER PAGE
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

# Project Title
title_para = doc.add_paragraph("F.R.I.E.N.D.S Fan Hub", style='CustomTitle')
title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle_para = doc.add_paragraph("ASP.NET Core MVC & .NET MAUI Mobile Application", style='CustomTitle')
subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_para.runs[0].font.size = Pt(18)

doc.add_paragraph("")
doc.add_paragraph("")

# Group Information
group_info = doc.add_paragraph("Group Number: ", style='CustomBody')
group_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
group_info.runs[0].font.bold = True
group_info.runs[0].font.size = Pt(14)
group_info.add_run("[AUTHOR]").font.size = Pt(14)

doc.add_paragraph("")

# Team Members Header
team_header = doc.add_paragraph("Team Members", style='CustomBody')
team_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
team_header.runs[0].font.bold = True
team_header.runs[0].font.size = Pt(14)

doc.add_paragraph("")

# Team Members Table
members_table = doc.add_table(rows=6, cols=3)
members_table.style = 'Table Grid'
members_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Header row
header_cells = members_table.rows[0].cells
header_cells[0].text = 'Name'
header_cells[1].text = 'Matric No.'
header_cells[2].text = 'Section'

# Make header bold
for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add team member rows (placeholder)
for i in range(1, 6):
    cells = members_table.rows[i].cells
    cells[0].text = f'[Team Member {i}]'
    cells[1].text = '[Matric No.]'
    cells[2].text = '[Section]'
    for cell in cells:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Set column widths
for row in members_table.rows:
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(2)
    row.cells[2].width = Inches(1.5)

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

# Date
date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %Y')}", style='CustomBody')
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
date_para.runs[0].font.size = Pt(12)

# Page break
doc.add_page_break()

# INTRODUCTION
intro_heading = doc.add_paragraph("1. INTRODUCTION", style='CustomHeading1')

intro_text = doc.add_paragraph(
    "The F.R.I.E.N.D.S Fan Hub is a comprehensive web and mobile application designed to celebrate "
    "and organize content from the beloved television series 'Friends'. This project demonstrates "
    "the implementation of a full-stack application using modern .NET technologies, showcasing both "
    "web and mobile development capabilities.", style='CustomBody'
)

doc.add_paragraph(
    "The project consists of two main components:", style='CustomBody'
)

bullet1 = doc.add_paragraph(style='CustomBody')
bullet1.add_run("• ").bold = True
bullet1.add_run("ASP.NET Core MVC Web Application: A feature-rich web application providing full CRUD "
                "operations for managing Friends-related content including characters, episodes, locations, and quotes.")

bullet2 = doc.add_paragraph(style='CustomBody')
bullet2.add_run("• ").bold = True
bullet2.add_run(".NET MAUI Mobile Application: A cross-platform mobile application that mirrors the "
                "functionality of the web application, providing users with access to the same content "
                "and features on mobile devices.")

doc.add_paragraph(
    "Both applications share a common SQL Server database, ensuring data consistency and demonstrating "
    "the ability to create unified solutions across different platforms. The project showcases modern "
    "development practices including Entity Framework Core for data access, responsive UI design, and "
    "proper separation of concerns through the MVC pattern.", style='CustomBody'
)

# Page break
doc.add_page_break()

# MODULES AND SUBMODULES
modules_heading = doc.add_paragraph("2. PROPOSED SYSTEM MODULES AND SUBMODULES", style='CustomHeading1')

# Module 1: Characters
mod1_heading = doc.add_paragraph("2.1 Characters Module", style='CustomHeading2')
doc.add_paragraph(
    "The Characters module manages information about the main and supporting characters from the Friends series.",
    style='CustomBody'
)

doc.add_paragraph("Submodules:", style='CustomBody').runs[0].bold = True
char_sub1 = doc.add_paragraph(style='CustomBody')
char_sub1.add_run("• Character List View: ").bold = True
char_sub1.add_run("Displays all characters with their images, names, and actor information")

char_sub2 = doc.add_paragraph(style='CustomBody')
char_sub2.add_run("• Character Details: ").bold = True
char_sub2.add_run("Shows comprehensive information including occupation, description, and media links")

char_sub3 = doc.add_paragraph(style='CustomBody')
char_sub3.add_run("• Create/Edit Character: ").bold = True
char_sub3.add_run("Forms for adding new characters or updating existing ones with image upload capability")

char_sub4 = doc.add_paragraph(style='CustomBody')
char_sub4.add_run("• Delete Character: ").bold = True
char_sub4.add_run("Confirmation-based deletion with proper validation")

# Module 2: Episodes
mod2_heading = doc.add_paragraph("2.2 Episodes Module", style='CustomHeading2')
doc.add_paragraph(
    "The Episodes module provides comprehensive episode management including season organization and media content.",
    style='CustomBody'
)

doc.add_paragraph("Submodules:", style='CustomBody').runs[0].bold = True
ep_sub1 = doc.add_paragraph(style='CustomBody')
ep_sub1.add_run("• Episode List: ").bold = True
ep_sub1.add_run("Season-wise episode listing with titles, air dates, and thumbnails")

ep_sub2 = doc.add_paragraph(style='CustomBody')
ep_sub2.add_run("• Episode Details: ").bold = True
ep_sub2.add_run("Full episode information including plot summary and video links")

ep_sub3 = doc.add_paragraph(style='CustomBody')
ep_sub3.add_run("• Episode Management: ").bold = True
ep_sub3.add_run("CRUD operations for episode data with media URL management")

ep_sub4 = doc.add_paragraph(style='CustomBody')
ep_sub4.add_run("• Season Navigation: ").bold = True
ep_sub4.add_run("Easy navigation between different seasons")

# Module 3: Locations
mod3_heading = doc.add_paragraph("2.3 Locations Module", style='CustomHeading2')
doc.add_paragraph(
    "The Locations module manages iconic locations from the series such as Central Perk, Monica's apartment, etc.",
    style='CustomBody'
)

doc.add_paragraph("Submodules:", style='CustomBody').runs[0].bold = True
loc_sub1 = doc.add_paragraph(style='CustomBody')
loc_sub1.add_run("• Location Gallery: ").bold = True
loc_sub1.add_run("Visual gallery of all locations with images")

loc_sub2 = doc.add_paragraph(style='CustomBody')
loc_sub2.add_run("• Location Information: ").bold = True
loc_sub2.add_run("Detailed descriptions and significance of each location")

loc_sub3 = doc.add_paragraph(style='CustomBody')
loc_sub3.add_run("• Location Management: ").bold = True
loc_sub3.add_run("Add, edit, and delete locations with image management")

loc_sub4 = doc.add_paragraph(style='CustomBody')
loc_sub4.add_run("• Location Types: ").bold = True
loc_sub4.add_run("Categorization of locations (apartments, workplaces, hangout spots)")

# Module 4: Quotes
mod4_heading = doc.add_paragraph("2.4 Quotes Module", style='CustomHeading2')
doc.add_paragraph(
    "The Quotes module captures memorable quotes and dialogues from the series, linked to characters and episodes.",
    style='CustomBody'
)

doc.add_paragraph("Submodules:", style='CustomBody').runs[0].bold = True
quote_sub1 = doc.add_paragraph(style='CustomBody')
quote_sub1.add_run("• Quote Browser: ").bold = True
quote_sub1.add_run("Browse quotes with character attribution and episode references")

quote_sub2 = doc.add_paragraph(style='CustomBody')
quote_sub2.add_run("• Quote Search: ").bold = True
quote_sub2.add_run("Search functionality to find specific quotes")

quote_sub3 = doc.add_paragraph(style='CustomBody')
quote_sub3.add_run("• Quote Management: ").bold = True
quote_sub3.add_run("Full CRUD operations with character and episode linking")

quote_sub4 = doc.add_paragraph(style='CustomBody')
quote_sub4.add_run("• Context Information: ").bold = True
quote_sub4.add_run("Episode and scene context for each quote")

# Page break
doc.add_page_break()

# Technical Architecture
tech_heading = doc.add_paragraph("2.5 Technical Architecture", style='CustomHeading2')

doc.add_paragraph("Web Application (ASP.NET Core MVC):", style='CustomBody').runs[0].bold = True
web_tech1 = doc.add_paragraph(style='CustomBody')
web_tech1.add_run("• ").bold = True
web_tech1.add_run("Entity Framework Core with SQL Server for data persistence")

web_tech2 = doc.add_paragraph(style='CustomBody')
web_tech2.add_run("• ").bold = True
web_tech2.add_run("Bootstrap 5 with custom CSS for responsive UI")

web_tech3 = doc.add_paragraph(style='CustomBody')
web_tech3.add_run("• ").bold = True
web_tech3.add_run("File upload functionality for images with server-side storage")

web_tech4 = doc.add_paragraph(style='CustomBody')
web_tech4.add_run("• ").bold = True
web_tech4.add_run("IIS deployment capability with proper configuration")

doc.add_paragraph("")

doc.add_paragraph("Mobile Application (.NET MAUI):", style='CustomBody').runs[0].bold = True
mobile_tech1 = doc.add_paragraph(style='CustomBody')
mobile_tech1.add_run("• ").bold = True
mobile_tech1.add_run("Cross-platform support (Android, iOS, Windows)")

mobile_tech2 = doc.add_paragraph(style='CustomBody')
mobile_tech2.add_run("• ").bold = True
mobile_tech2.add_run("MVVM architecture with data binding")

mobile_tech3 = doc.add_paragraph(style='CustomBody')
mobile_tech3.add_run("• ").bold = True
mobile_tech3.add_run("Shared database access with the web application")

mobile_tech4 = doc.add_paragraph(style='CustomBody')
mobile_tech4.add_run("• ").bold = True
mobile_tech4.add_run("Custom UI with glassmorphism design and neon accents")

# Page break
doc.add_page_break()

# DISCUSSION AND CONCLUSION
discussion_heading = doc.add_paragraph("3. DISCUSSION AND CONCLUSION", style='CustomHeading1')

doc.add_paragraph("3.1 Project Achievements", style='CustomHeading2')

doc.add_paragraph(
    "This project successfully demonstrates the implementation of a comprehensive fan hub application "
    "using modern .NET technologies. Key achievements include:",
    style='CustomBody'
)

achieve1 = doc.add_paragraph(style='CustomBody')
achieve1.add_run("• ").bold = True
achieve1.add_run("Successful implementation of all four required modules with complete CRUD operations")

achieve2 = doc.add_paragraph(style='CustomBody')
achieve2.add_run("• ").bold = True
achieve2.add_run("Seamless data sharing between web and mobile applications through a unified database")

achieve3 = doc.add_paragraph(style='CustomBody')
achieve3.add_run("• ").bold = True
achieve3.add_run("Responsive and user-friendly interfaces on both platforms")

achieve4 = doc.add_paragraph(style='CustomBody')
achieve4.add_run("• ").bold = True
achieve4.add_run("Implementation of advanced features like file uploads and media management")

doc.add_paragraph("3.2 Technical Challenges and Solutions", style='CustomHeading2')

doc.add_paragraph(
    "During development, several technical challenges were encountered and successfully resolved:",
    style='CustomBody'
)

challenge1 = doc.add_paragraph(style='CustomBody')
challenge1.add_run("• ").bold = True
challenge1.add_run("Database Connectivity: ")
challenge1.add_run("Configuring the MAUI application to connect to the same SQL Server database "
                   "required careful handling of connection strings and network permissions, especially "
                   "for Android emulator connectivity.")

challenge2 = doc.add_paragraph(style='CustomBody')
challenge2.add_run("• ").bold = True
challenge2.add_run("Image Management: ")
challenge2.add_run("Implementing a unified image management system that works across both web and "
                   "mobile platforms required creating a flexible URL resolution system with fallback "
                   "mechanisms for missing images.")

challenge3 = doc.add_paragraph(style='CustomBody')
challenge3.add_run("• ").bold = True
challenge3.add_run("UI Consistency: ")
challenge3.add_run("Maintaining visual consistency between the web and mobile applications while "
                   "respecting platform-specific design guidelines required careful planning and "
                   "custom styling implementation.")

doc.add_paragraph("3.3 Learning Outcomes", style='CustomHeading2')

doc.add_paragraph(
    "This project provided valuable learning experiences in several areas:",
    style='CustomBody'
)

learning1 = doc.add_paragraph(style='CustomBody')
learning1.add_run("• ").bold = True
learning1.add_run("Full-stack .NET development across multiple platforms")

learning2 = doc.add_paragraph(style='CustomBody')
learning2.add_run("• ").bold = True
learning2.add_run("Database design and Entity Framework Core implementation")

learning3 = doc.add_paragraph(style='CustomBody')
learning3.add_run("• ").bold = True
learning3.add_run("Cross-platform mobile development with .NET MAUI")

learning4 = doc.add_paragraph(style='CustomBody')
learning4.add_run("• ").bold = True
learning4.add_run("Project organization and code reusability strategies")

doc.add_paragraph("3.4 Future Enhancements", style='CustomHeading2')

doc.add_paragraph(
    "While the project meets all requirements, several enhancements could be implemented:",
    style='CustomBody'
)

future1 = doc.add_paragraph(style='CustomBody')
future1.add_run("• ").bold = True
future1.add_run("User authentication and personalized content")

future2 = doc.add_paragraph(style='CustomBody')
future2.add_run("• ").bold = True
future2.add_run("Social features like favorite quotes and character ratings")

future3 = doc.add_paragraph(style='CustomBody')
future3.add_run("• ").bold = True
future3.add_run("Advanced search and filtering capabilities")

future4 = doc.add_paragraph(style='CustomBody')
future4.add_run("• ").bold = True
future4.add_run("Integration with external APIs for additional content")

doc.add_paragraph("3.5 Conclusion", style='CustomHeading2')

doc.add_paragraph(
    "The F.R.I.E.N.D.S Fan Hub project successfully demonstrates the ability to develop a comprehensive, "
    "multi-platform application using modern .NET technologies. By implementing both ASP.NET Core MVC "
    "for web and .NET MAUI for mobile, the project showcases the versatility and power of the .NET "
    "ecosystem. The successful implementation of all four required modules with complete CRUD operations, "
    "along with the seamless data sharing between platforms, validates the architectural decisions made "
    "during development.",
    style='CustomBody'
)

doc.add_paragraph(
    "This project not only meets all the specified requirements but also demonstrates best practices in "
    "software development including proper separation of concerns, responsive design, and efficient data "
    "management. The experience gained from this project provides a solid foundation for future development "
    "in both web and mobile application domains using .NET technologies.",
    style='CustomBody'
)

# Save the document
doc.save('FriendsFanHub_Project_Report.docx')
print("Project report generated successfully as 'FriendsFanHub_Project_Report.docx'")